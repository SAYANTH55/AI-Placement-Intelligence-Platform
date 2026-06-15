import os
import sys
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Define paths
FONT_PATH = "C:\\Windows\\Fonts\\arial.ttf"
BOLD_FONT_PATH = "C:\\Windows\\Fonts\\arialbd.ttf"
DIAGRAM_PATH = os.path.join(os.path.dirname(__file__), "workflow_diagram.png")
DOCX_PATH = os.path.join(os.path.dirname(__file__), "COPYRIGHT_APPLICATION_DRAFT_SPREAD.docx")

def generate_diagram():
    """
    Generates a highly professional high-resolution PNG image of the workflow
    using Pillow to insert as a diagram inside the Word document.
    Horizontal 2-row layout with wide margins to prevent text crowding.
    """
    print("Generating workflow diagram image...")
    # Create wide high-res canvas
    width, height = 1300, 460
    img = Image.new("RGB", (width, height), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    
    # Load fonts
    try:
        font_title = ImageFont.truetype(BOLD_FONT_PATH, 19)
        font_step_num = ImageFont.truetype(BOLD_FONT_PATH, 14)
        font_step_title = ImageFont.truetype(BOLD_FONT_PATH, 12)
        font_step_desc = ImageFont.truetype(FONT_PATH, 9.5)
        font_connector = ImageFont.truetype(BOLD_FONT_PATH, 9.5) # Bold connectors for readability
    except IOError:
        font_title = font_step_num = font_step_title = font_step_desc = font_connector = ImageFont.load_default()

    # Draw Title
    title_text = "AI-Driven Placement Intelligence Platform - Process Pipeline & Data Workflow"
    draw.text((width // 2, 28), title_text, fill="#111827", font=font_title, anchor="mm")
    
    # Define node helper
    def draw_node(x, y, w, h, step_num, title, desc_lines, fill_color, border_color, text_color, title_color):
        # Draw shadow
        draw.rounded_rectangle([x+3, y+3, x+w+3, y+h+3], radius=6, fill="#E2E8F0")
        # Draw background and border
        draw.rounded_rectangle([x, y, x+w, y+h], radius=6, fill=fill_color, outline=border_color, width=2)
        # Draw step number
        draw.rounded_rectangle([x + 10, y + 10, x + 34, y + 32], radius=4, fill=border_color)
        draw.text((x + 22, y + 21), step_num, fill="#FFFFFF", font=font_step_num, anchor="mm")
        # Draw title
        draw.text((x + 44, y + 21), title, fill=title_color, font=font_step_title, anchor="lm")
        # Draw description lines
        curr_y = y + 44
        for line in desc_lines:
            draw.text((x + 12, curr_y), line, fill=text_color, font=font_step_desc)
            curr_y += 16

    # Dashed line helper for the regression tracking feedback loop
    def draw_dashed_line(start, end, dash_length=6, gap_length=4, fill="#94A3B8", width=2):
        x1, y1 = start
        x2, y2 = end
        dx = x2 - x1
        dy = y2 - y1
        dist = (dx**2 + dy**2)**0.5
        if dist == 0:
            return
        dx /= dist
        dy /= dist
        
        curr_len = 0
        draw_state = True
        while curr_len < dist:
            step = dash_length if draw_state else gap_length
            if curr_len + step > dist:
                step = dist - curr_len
            if draw_state:
                draw.line([(x1 + curr_len*dx, y1 + curr_len*dy), 
                           (x1 + (curr_len+step)*dx, y1 + (curr_len+step)*dy)], 
                          fill=fill, width=width)
            curr_len += step
            draw_state = not draw_state

    # Top Row Y = 75, Height = 115, Width = 230
    y_top = 75
    h_top = 115
    w_node = 230

    # Bottom Row Y = 265, Height = 135, Width = 230
    y_bottom = 265
    h_bottom = 135

    # Node X-positions (Gaps are 95px wide: 45 -> 275 -> 370 -> 600 -> 695 -> 925 -> 1020 -> 1250)
    x0 = 45
    x1 = 370
    x2 = 695
    x3 = 1020

    # Steps Configuration
    # Step 1: Input Ingestion
    draw_node(x0, y_top, w_node, h_top, 
              "01", "Input & Ingestion", 
              ["* React 18 Dropzone interface", "* File types: .pdf, .docx, .txt", "* Dynamic transfer via Axios"], 
              "#EFF6FF", "#3B82F6", "#1E3A8A", "#1E3A8A")

    # Step 2: FastAPI Handling
    draw_node(x1, y_top, w_node, h_top, 
              "02", "FastAPI Endpoints", 
              ["* Receives multipart/form-data", "* Async route execution", "* Secures temp file streams"], 
              "#F5F3FF", "#8B5CF6", "#3730A3", "#3730A3")

    # Step 3: Text Extraction
    draw_node(x2, y_top, w_node, h_top, 
              "03", "Text Extraction", 
              ["* pdfplumber (for PDF formats)", "* python-docx (for Word formats)", "* Automatic temp file cleanup"], 
              "#ECFDF5", "#10B981", "#064E3B", "#064E3B")

    # Step 4: NLP & Fuzzy Normalizer
    draw_node(x3, y_top, w_node, h_top, 
              "04", "Fuzzy Normalization", 
              ["* spaCy custom tokenization", "* Normalized via Levenshtein (&ge;80%)", "* Standard catalog mapping", "* Regex experience extraction"], 
              "#FEF3C7", "#F59E0B", "#78350F", "#78350F")

    # Step 5: Dual Evaluation
    draw_node(x3, y_bottom, w_node, h_bottom, 
              "05", "Core Intelligence", 
              ["* A. Weighted Job Matcher", "  - Core vs. Peripheral skills", "* B. RandomForest ML model", "  - Predicts placement readiness", "  - GPA, skills, exp, projects"], 
              "#FFEDD5", "#F97316", "#7C2D12", "#7C2D12")

    # Step 6: Relational Persistence
    draw_node(x2, y_bottom, w_node, h_bottom, 
              "06", "Relational Persistence", 
              ["* Persisted via SQLAlchemy ORM", "* SQLite 'resume_uploads' logs", "* User profile analytics snapshot", "* Upload progress tracker"], 
              "#F0FDF4", "#10B981", "#064E3B", "#064E3B")

    # Step 7: JSON Delivery
    draw_node(x1, y_bottom, w_node, h_bottom, 
              "07", "Async Delivery", 
              ["* Generates secure JSON payload", "* Includes matched skill gaps,", "  readiness levels, and customized", "  learning recommendations"], 
              "#FAF5FF", "#A855F7", "#581C87", "#581C87")

    # Step 8: Dashboard Re-rendering
    draw_node(x0, y_bottom, w_node, h_bottom, 
              "08", "Interactive Display", 
              ["* Framer Motion animations", "* Interactive score rings", "* Skill analysis segment lists", "* Reportlab PDF exporter"], 
              "#ECFEFF", "#06B6D4", "#083344", "#083344")

    # DRAW CONNECTORS (ARROWS)
    def draw_arrow(start_pos, end_pos, label="", direction="right"):
        draw.line([start_pos, end_pos], fill="#9CA3AF", width=2)
        # Draw Arrow head
        x1, y1 = start_pos
        x2, y2 = end_pos
        if direction == "down":
            draw.polygon([(x2-5, y2-7), (x2+5, y2-7), (x2, y2)], fill="#4B5563")
        elif direction == "right":
            draw.polygon([(x2-7, y2-5), (x2-7, y2+5), (x2, y2)], fill="#4B5563")
        elif direction == "left":
            draw.polygon([(x2+7, y2-5), (x2+7, y2+5), (x2, y2)], fill="#4B5563")
            
        if label:
            lx = (x1 + x2) // 2
            ly = ((y1 + y2) // 2) - 10
            draw.text((lx, ly), label, fill="#4B5563", font=font_connector, anchor="mm")

    # Step 1 -> 2
    draw_arrow((x0 + w_node, y_top + h_top // 2), (x1, y_top + h_top // 2), "FormData", "right")
    # Step 2 -> 3
    draw_arrow((x1 + w_node, y_top + h_top // 2), (x2, y_top + h_top // 2), "Temp File", "right")
    # Step 3 -> 4
    draw_arrow((x2 + w_node, y_top + h_top // 2), (x3, y_top + h_top // 2), "Raw Text", "right")
    
    # Step 4 -> 5 (Vertical Connect Down)
    draw_arrow((x3 + w_node // 2, y_top + h_top), (x3 + w_node // 2, y_bottom), "Norm Skills", "down")
    
    # Step 5 -> 6
    draw_arrow((x3, y_bottom + h_bottom // 2), (x2 + w_node, y_bottom + h_bottom // 2), "Profiles", "left")
    # Step 6 -> 7
    draw_arrow((x2, y_bottom + h_bottom // 2), (x1 + w_node, y_bottom + h_bottom // 2), "Save Success", "left")
    # Step 7 -> 8
    draw_arrow((x1, y_bottom + h_bottom // 2), (x0 + w_node, y_bottom + h_bottom // 2), "JSON Payload", "left")

    # Feedback Loop: Step 8 -> 1 (Vertical Connect Up, Dashed)
    loop_start = (x0 + w_node // 2, y_bottom)
    loop_end = (x0 + w_node // 2, y_top + h_top)
    draw_dashed_line(loop_start, (loop_end[0], loop_end[1] + 10), fill="#94A3B8", width=2)
    # Draw arrow head pointing up at Step 1 bottom
    draw.polygon([(loop_end[0]-5, loop_end[1]+10), (loop_end[0]+5, loop_end[1]+10), loop_end], fill="#94A3B8")
    # Label the feedback loop (Drawn off-center to prevent any vertical text crowding)
    draw.text((x0 + w_node // 2 + 10, (y_top + h_top + y_bottom) // 2), "Re-Upload Cycle", fill="#94A3B8", font=font_connector, anchor="lm")

    # Save image
    img.save(DIAGRAM_PATH, "PNG")
    print(f"Diagram image saved successfully at: {DIAGRAM_PATH}")

def set_cell_background(cell, color_hex):
    """Sets background color of a Word table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def create_word_document():
    """
    Builds the professional word document (.docx) embedding the generated diagram.
    """
    print("Building MS Word document...")
    doc = Document()

    # Set Margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base typography (Arial)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700

    # Header / Title
    title = doc.add_paragraph()
    title_run = title.add_run("COPYRIGHT REGISTRATION SPECIFICATION")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Slate 900
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Technical Document of Original Software Authorship")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray 600
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)

    # Metadata Table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Light Shading Accent 1'
    table.autofit = False
    
    # Format cells
    col_widths = [Inches(2.2), Inches(4.3)]
    row_data = [
        ("Title of the Work", "AI-Driven Placement Intelligence Platform (v2.0.0)"),
        ("Class of Work", "Computer Software / Literary Work (Software Source Code)")
    ]

    for i, row in enumerate(table.rows):
        # Set height
        row.height = Pt(24)
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]
            cell.text = row_data[i][j]
            
            # Apply padding and styling
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            if j == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                set_cell_background(cell, "F1F5F9")
            else:
                run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                set_cell_background(cell, "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Section Helper
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C) # Orange 600

    # Section 1: Problem
    add_section_heading("1. Statement of the Problem")
    
    p = doc.add_paragraph("In the modern academic and corporate recruitment landscape, graduating students and higher education institutions face several systemic challenges regarding career readiness and placement optimization:")
    p.paragraph_format.space_after = Pt(10)

    problems = [
        ("Unstructured Resume Data Processing", "Resumes are authored in varying, unstructured formats (PDF, DOCX, TXT) with highly diverse semantic phrasing. Traditional automated screening systems (Applicant Tracking Systems or ATS) rely on exact keyword matches, leading to high failure rates when identifying qualified candidates who use non-standard terminology or abbreviations (e.g., writing 'js' instead of 'JavaScript', or 'sklearn' instead of 'scikit-learn')."),
        ("Opaque Skill Gap Identification", "Students frequently lack an objective, detailed understanding of how their current competencies align with specific industry roles (e.g., Frontend Developer, Backend Developer, DevOps, Data Science). There is no automated, weighted mechanism available to differentiate between 'Core' skills (indispensable for a role) and 'Peripheral' skills, leading to inefficient self-guided learning."),
        ("Absence of Objective Placement Readiness Metrics", "Career placement cells and students have no predictive, data-driven tools to objectively quantify placement probability. Placement readiness is historically evaluated through subjective human reviews or raw academic GPAs, which fail to capture the multi-dimensional synergy between GPA, total technical skill diversity, project counts, and prior internship experience."),
        ("Administrative Overhead in Progress Tracking", "Placement officers and student advisors are forced to manually track the progress of students over multiple resume iterations, lacking a centralized, secure repository that logs historical resume analyses, rate-limits abuse, and automatically generates actionable, dynamic study recommendations.")
    ]

    for title, desc in problems:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(6)
        r_title = bp.add_run(f"{title}: ")
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        bp.add_run(desc)

    # Section 2: Solution
    add_section_heading("2. Proposed Solution")
    
    p = doc.add_paragraph("The AI-Driven Placement Intelligence Platform is a proprietary, multi-tiered software application designed to solve these challenges by introducing an automated, end-to-end intelligence layer. The software represents an original integration of modern web technologies, Natural Language Processing (NLP), and machine learning pipelines:")
    p.paragraph_format.space_after = Pt(10)

    solutions = [
        ("Intelligent Text Ingestion & Parsing", "The platform supports cross-format resume ingestion (PDF, DOCX, TXT) and features an asynchronous text extraction engine. It leverages advanced NLP libraries to parse raw text streams and immediately executes secure memory clean-ups to guarantee candidate data privacy."),
        ("NLP-Powered Fuzzy Skill Normalization", "To resolve the keyword-matching bottleneck, the platform implements a specialized Fuzzy Skill Normalizer. It compares raw resume terms against a standardized industry catalog using a similarity threshold (set at 80%). This normalizes highly diverse student inputs into canonical industry skills, ensuring that technical capability is recognized regardless of stylistic phrasing."),
        ("Dynamic Weighted Multi-Role Matcher", "The platform features a customized multi-role scoring engine. Instead of simple keyword counts, it evaluates skills across various professional roles, classifying them into a weighted hierarchy (Core vs. Peripheral). This yields an exact match percentage, present competencies, and an actionable Skill Gap list detailing the precise technologies a student needs to master."),
        ("Predictive Machine Learning Engine", "The software incorporates a serialized Random Forest Classifier model trained on historical student placement profiles. By creating a four-dimensional feature vector—comprising cumulative GPA, total technical skills, academic/industry projects, and years of work experience—the model outputs an objective, mathematical probability of placement and maps the user to an ordinal readiness tier (High, Medium, or Low)."),
        ("Secure, Distributed Architecture", "The system employs a secure three-tier design. This consists of a React 18 & Vite Frontend Portal, FastAPI middleware featuring bcrypt hashing and Brevo OTP verifications, and a relational database layer backed by SQLite and SQLAlchemy ORM for detailed analytics snapshot persistence.")
    ]

    for title, desc in solutions:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(6)
        r_title = bp.add_run(f"{title}: ")
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        bp.add_run(desc)

    # Section 3: Workflow
    add_section_heading("3. Software Workflow & Data Pipeline")
    
    p = doc.add_paragraph("The workflow of the software maps the transition of unstructured input files into highly structured, actionable prediction insights. Below is the technical flow of the application:")
    p.paragraph_format.space_after = Pt(14)

    # Insert Workflow Diagram
    if os.path.exists(DIAGRAM_PATH):
        print("Inserting workflow diagram to document...")
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(DIAGRAM_PATH, width=Inches(6.2))
        pic_p.paragraph_format.space_after = Pt(10)
        
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run("Figure 1: Architectural Data Pipeline and Processing Flowchart")
        caption_run.font.italic = True
        caption_run.font.size = Pt(9.5)
        caption_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        caption_p.paragraph_format.space_after = Pt(18)
    else:
        print(f"Warning: Diagram image not found at {DIAGRAM_PATH}")

    # Describe stages
    stages = [
        ("Input & Ingestion", "The frontend application validates file constraints and transmits the resume via Axios in multipart/form-data encoding to /upload_resume."),
        ("Asynchronous Parsing", "FastAPI captures the upload, creates a temporary file, detects the MIME-type, and utilizes pdfplumber or python-docx to convert document buffers into raw text strings. The temp file is purged instantly upon text extraction."),
        ("NLP Skill Normalization", "The text is cleaned and evaluated against a precompiled SKILLS_DICTIONARY using the Fuzzy Skill Normalizer (skill_normalizer.py). Matches having an 80% or greater Levenshtein similarity are canonicalized to standard terms. Regular expression matching extracts years of experience."),
        ("Joint Evaluation Pipeline", "The extracted skills and attributes are parsed through a dual evaluation core: (A) Weighted Job Matcher distinguishes Core vs. Peripheral skills; (B) Random Forest Readiness Classifier predicts a probabilistic placement percentage from 0% to 100% based on feature inputs."),
        ("Relational Logging & Presentation", "The results are serialized as JSON and logged inside the SQLite resume_uploads table. The backend returns a structured JSON payload to React which updates its state context and triggers an animated dashboard rendering with interactive score rings, and dynamic learning tips. Users can trigger custom print-ready Reportlab PDF summaries.")
    ]

    for i, (title, desc) in enumerate(stages):
        sp = doc.add_paragraph()
        sp.paragraph_format.left_indent = Inches(0.25)
        sp.paragraph_format.space_after = Pt(6)
        r_num = sp.add_run(f"{i+1}. {title}: ")
        r_num.bold = True
        r_num.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        sp.add_run(desc)

    # Section 4: Conclusion
    add_section_heading("4. Conclusion & Claims of Originality")
    
    p1 = doc.add_paragraph("The AI-Driven Placement Intelligence Platform is a fully integrated, proprietary software solution that systematically addresses structural friction in modern career placement and student preparedness scoring. By coupling high-throughput asynchronous execution, specialized natural language processing (NLP) fuzzy-mapping models, and advanced random forest classifiers, the software resolves the limitations of standard screening procedures.")
    p1.paragraph_format.space_after = Pt(10)

    p2 = doc.add_paragraph("The software is an original literary and technical creation of the authors. Its custom-engineered components—specifically the weighted core/peripheral skill gap matcher, the automated fuzzy skill normalizer, the dynamic predictive feature model, and the unified student-focused analytics interface—are unique, functional creations worthy of copyright protection. Copyright registration will protect the entire source code base, database schema layouts, custom integrated algorithmic routines, and dynamic web interfaces from unauthorized replication, reproduction, or deployment.")
    p2.paragraph_format.space_after = Pt(18)

    # Save
    doc.save(DOCX_PATH)
    print(f"MS Word document built and saved at: {DOCX_PATH}")

if __name__ == "__main__":
    generate_diagram()
    create_word_document()
    print("Done compiling copyright package!")
